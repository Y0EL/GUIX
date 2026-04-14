from __future__ import annotations

import csv
import hashlib
import json
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List
from xml.sax.saxutils import escape as escape_xml

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from orchestration.config import PengaturanRuntime
from orchestration.logging_utils import logger_dengan_trace
from orchestration.mcp import McpGateway
from orchestration.openai_stack import LayananOpenAI
from orchestration.schema import DossierSindikat


WARNA_UTAMA = "E8682A"
WARNA_SEKUNDER = "F5A05A"
WARNA_LATAR = "FFFFFF"
WARNA_PERMUKAAN = "FFF8F5"
WARNA_TEKS = "1A1A1A"
WARNA_TEKS_REDUP = "6B6B6B"
WARNA_BATAS = "E5D5CC"
WARNA_SUKSES = "2E7D32"
WARNA_PERINGATAN = "E65100"
WARNA_BAHAYA = "C62828"
NAMA_FONT_UI = "Segoe UI"


def _slug(teks: str) -> str:
    hasil = []
    for karakter in teks.lower():
        if karakter.isalnum():
            hasil.append(karakter)
        elif hasil and hasil[-1] != "-":
            hasil.append("-")
    return "".join(hasil).strip("-") or "analisa"


def _json_stabil(payload: Any) -> str:
    def _serializer_default(nilai: Any) -> Any:
        if isinstance(nilai, (datetime, date)):
            return nilai.isoformat()
        if isinstance(nilai, Path):
            return str(nilai)
        return str(nilai)

    return json.dumps(
        payload,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=_serializer_default,
    )


def _json_rapi(payload: Any) -> str:
    def _serializer_default(nilai: Any) -> Any:
        if isinstance(nilai, (datetime, date)):
            return nilai.isoformat()
        if isinstance(nilai, Path):
            return str(nilai)
        return str(nilai)

    return json.dumps(
        payload,
        ensure_ascii=False,
        indent=2,
        default=_serializer_default,
    )


def _sha256_teks(teks: str) -> str:
    return hashlib.sha256(teks.encode("utf-8")).hexdigest()


class PenghasilArtefakAnalisa:
    def __init__(self, akar_output: Path):
        self.akar_output = akar_output

    def simpan_semua(
        self,
        id_kasus: str,
        dossier: DossierSindikat,
        bundel: Dict[str, Any],
    ) -> Dict[str, Any]:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        payload_sidik = {
            "dossier": dossier.model_dump(mode="json"),
            "bundel": bundel,
            "dibuat_pada": timestamp,
        }
        sidik = _sha256_teks(_json_stabil(payload_sidik))
        sidik_pendek = sidik[:16]
        nama_dasar = f"{_slug(id_kasus)}_{timestamp}_{sidik_pendek}"
        direktori = self.akar_output / _slug(id_kasus) / nama_dasar
        direktori.mkdir(parents=True, exist_ok=True)

        jalur_json = direktori / f"{nama_dasar}.json"
        jalur_md = direktori / f"{nama_dasar}.md"
        jalur_ringkasan_csv = direktori / f"{nama_dasar}_ringkasan.csv"
        jalur_aktor_csv = direktori / f"{nama_dasar}_aktor.csv"
        jalur_relasi_csv = direktori / f"{nama_dasar}_relasi.csv"
        jalur_transaksi_csv = direktori / f"{nama_dasar}_transaksi.csv"
        jalur_xlsx = direktori / f"{nama_dasar}.xlsx"
        jalur_docx = direktori / f"{nama_dasar}.docx"
        jalur_pdf = direktori / f"{nama_dasar}.pdf"
        jalur_manifest = direktori / "manifest.json"

        isi_md = self._bangun_markdown(dossier, bundel, sidik)
        payload_json = {
            "fingerprint_sha256": sidik,
            "dibuat_pada": timestamp,
            "id_kasus": id_kasus,
            "dossier": dossier.model_dump(mode="json"),
            "bundel_ringkas": self._ringkas_bundel(bundel),
            "bundel_lengkap": bundel,
        }

        jalur_json.write_text(_json_rapi(payload_json), encoding="utf-8")
        jalur_md.write_text(isi_md, encoding="utf-8")

        self._tulis_csv_ringkasan(jalur_ringkasan_csv, dossier, sidik)
        self._tulis_csv_aktor(jalur_aktor_csv, dossier)
        self._tulis_csv_relasi(jalur_relasi_csv, dossier)
        self._tulis_csv_transaksi(jalur_transaksi_csv, bundel.get("transaksi", []))
        self._tulis_xlsx(jalur_xlsx, dossier, bundel, sidik)
        self._tulis_docx(jalur_docx, dossier, bundel, sidik)
        self._tulis_pdf(jalur_pdf, dossier, bundel, sidik)

        manifest = {
            "id_kasus": id_kasus,
            "dibuat_pada": timestamp,
            "fingerprint_sha256": sidik,
            "artefak": [],
        }
        for jalur in [
            jalur_json,
            jalur_md,
            jalur_ringkasan_csv,
            jalur_aktor_csv,
            jalur_relasi_csv,
            jalur_transaksi_csv,
            jalur_xlsx,
            jalur_docx,
            jalur_pdf,
        ]:
            manifest["artefak"].append(
                {
                    "nama_file": jalur.name,
                    "ukuran_byte": jalur.stat().st_size,
                    "sha256": _sha256_teks(jalur.read_text(encoding="utf-8", errors="ignore"))
                    if jalur.suffix.lower() in {".json", ".md", ".csv"}
                    else hashlib.sha256(jalur.read_bytes()).hexdigest(),
                }
            )
        jalur_manifest.write_text(_json_rapi(manifest), encoding="utf-8")
        return {
            "direktori": str(direktori),
            "fingerprint_sha256": sidik,
            "files": [item["nama_file"] for item in manifest["artefak"]] + [jalur_manifest.name],
        }

    def _ringkas_bundel(self, bundel: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "jumlah_laporan": len(bundel.get("laporan", [])),
            "jumlah_skor_risiko": len(bundel.get("skor_risiko", [])),
            "jumlah_transaksi": len(bundel.get("transaksi", [])),
            "jumlah_kampanye": len(bundel.get("kampanye", [])),
            "jumlah_profil": len(bundel.get("profil", [])),
            "jumlah_lokasi": len(bundel.get("lokasi", [])),
            "jumlah_postingan": len(bundel.get("postingan", [])),
            "jumlah_node_graf": len((bundel.get("graf") or {}).get("nodes", [])),
            "jumlah_edge_graf": len((bundel.get("graf") or {}).get("edges", [])),
        }

    def _bangun_markdown(self, dossier: DossierSindikat, bundel: Dict[str, Any], sidik: str) -> str:
        baris: List[str] = []
        baris.append(f"# Dossier Analisa Kasus - {dossier.judul_kasus}")
        baris.append("")
        baris.append(f"- ID kasus: `{dossier.id_kasus}`")
        baris.append(f"- Fingerprint SHA-256: `{sidik}`")
        baris.append(f"- Indikasi sindikat: `{'ya' if dossier.indikasi_sindikat else 'tidak'}`")
        baris.append(f"- Confidence: `{dossier.confidence}`")
        baris.append("")
        baris.append("## Ringkasan Eksekutif")
        baris.append("")
        baris.append(dossier.ringkasan_eksekutif)
        baris.append("")
        baris.append("## Alasan Utama")
        baris.append("")
        for item in dossier.alasan_utama:
            baris.append(f"- {item}")
        baris.append("")
        baris.append("## Pola Koordinasi")
        baris.append("")
        for item in dossier.pola_koordinasi:
            baris.append(f"- {item}")
        baris.append("")
        baris.append("## Aktor Inti")
        baris.append("")
        for item in dossier.aktor_inti:
            baris.append(
                f"- **{item.nama}** (`{item.id_profil or '-'}`) | peran: {item.peran} | confidence: {item.confidence} | alasan: {item.alasan}"
            )
        baris.append("")
        baris.append("## Relasi Kunci")
        baris.append("")
        for item in dossier.relasi_kunci:
            baris.append(
                f"- {item.sumber} -> {item.target} | {item.jenis_relasi} | confidence: {item.confidence} | alasan: {item.alasan}"
            )
        baris.append("")
        baris.append("## Bukti Utama")
        baris.append("")
        for item in dossier.bukti_utama:
            baris.append(f"- {item}")
        baris.append("")
        baris.append("## Bukti Lemah")
        baris.append("")
        for item in dossier.bukti_lemah:
            baris.append(f"- {item}")
        baris.append("")
        baris.append("## Rekomendasi Lanjutan")
        baris.append("")
        for item in dossier.rekomendasi_lanjutan:
            baris.append(f"- {item}")
        baris.append("")
        baris.append("## Narasi Analisis")
        baris.append("")
        baris.append(dossier.narasi_analisis)
        baris.append("")
        baris.append("## Statistik Bundel")
        baris.append("")
        for kunci, nilai in self._ringkas_bundel(bundel).items():
            baris.append(f"- {kunci}: {nilai}")
        baris.append("")
        return "\n".join(baris)

    def _tulis_csv_ringkasan(self, jalur: Path, dossier: DossierSindikat, sidik: str) -> None:
        with jalur.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["kunci", "nilai"])
            writer.writerow(["id_kasus", dossier.id_kasus])
            writer.writerow(["judul_kasus", dossier.judul_kasus])
            writer.writerow(["indikasi_sindikat", "ya" if dossier.indikasi_sindikat else "tidak"])
            writer.writerow(["confidence", dossier.confidence])
            writer.writerow(["fingerprint_sha256", sidik])
            writer.writerow(["confidence_reasoning", dossier.confidence_reasoning])

    def _tulis_csv_aktor(self, jalur: Path, dossier: DossierSindikat) -> None:
        with jalur.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["id_profil", "nama", "peran", "confidence", "alasan"])
            for item in dossier.aktor_inti:
                writer.writerow([item.id_profil, item.nama, item.peran, item.confidence, item.alasan])

    def _tulis_csv_relasi(self, jalur: Path, dossier: DossierSindikat) -> None:
        with jalur.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["sumber", "target", "jenis_relasi", "confidence", "alasan"])
            for item in dossier.relasi_kunci:
                writer.writerow([item.sumber, item.target, item.jenis_relasi, item.confidence, item.alasan])

    def _tulis_csv_transaksi(self, jalur: Path, transaksi: Iterable[Dict[str, Any]]) -> None:
        with jalur.open("w", encoding="utf-8", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(["id_transaksi", "nama_sumber", "nama_tujuan", "jumlah_idr", "kanal", "timestamp"])
            for item in transaksi:
                writer.writerow(
                    [
                        item.get("id_transaksi", ""),
                        item.get("nama_sumber", ""),
                        item.get("nama_tujuan", ""),
                        item.get("jumlah_idr", ""),
                        item.get("kanal", ""),
                        item.get("timestamp", ""),
                    ]
                )

    def _ya_tidak(self, nilai: bool) -> str:
        return "Ya" if nilai else "Tidak"

    def _format_persentase(self, nilai: Any) -> str:
        try:
            angka = float(nilai)
        except (TypeError, ValueError):
            return str(nilai)
        if 0 <= angka <= 1:
            angka *= 100
        return f"{angka:.0f}%"

    def _format_uang_rupiah(self, nilai: Any) -> str:
        try:
            angka = float(nilai)
        except (TypeError, ValueError):
            return str(nilai)
        return f"Rp {angka:,.0f}".replace(",", ".")

    def _format_waktu(self, nilai: Any) -> str:
        if isinstance(nilai, datetime):
            return nilai.isoformat(sep=" ", timespec="seconds")
        return str(nilai or "-")

    def _isi_sel_xlsx(self, sel, *, tebal: bool = False, warna_teks: str = WARNA_TEKS, warna_isi: str = WARNA_LATAR) -> None:
        sel.font = Font(name=NAMA_FONT_UI, size=11, bold=tebal, color=warna_teks)
        sel.fill = PatternFill(fill_type="solid", fgColor=warna_isi)
        sel.border = Border(
            left=Side(style="thin", color=WARNA_BATAS),
            right=Side(style="thin", color=WARNA_BATAS),
            top=Side(style="thin", color=WARNA_BATAS),
            bottom=Side(style="thin", color=WARNA_BATAS),
        )
        sel.alignment = Alignment(vertical="top", wrap_text=True)

    def _tambah_judul_seksi_xlsx(self, lembar, baris: int, judul: str, kolom_akhir: str) -> int:
        lembar.merge_cells(f"A{baris}:{kolom_akhir}{baris}")
        sel = lembar[f"A{baris}"]
        sel.value = judul
        self._isi_sel_xlsx(sel, tebal=True, warna_teks=WARNA_LATAR, warna_isi=WARNA_UTAMA)
        sel.alignment = Alignment(horizontal="left", vertical="center")
        return baris + 1

    def _atur_lebar_kolom_xlsx(self, lembar) -> None:
        for kolom in lembar.columns:
            indeks = kolom[0].column
            panjang = max(len(str(sel.value or "")) for sel in kolom)
            lembar.column_dimensions[get_column_letter(indeks)].width = min(max(panjang + 3, 12), 40)

    def _tulis_tabel_xlsx(self, lembar, baris_mulai: int, header: List[str], data: List[List[Any]], freeze_pane: str | None = None) -> None:
        for indeks, judul in enumerate(header, start=1):
            sel = lembar.cell(row=baris_mulai, column=indeks, value=judul)
            self._isi_sel_xlsx(sel, tebal=True, warna_teks=WARNA_LATAR, warna_isi=WARNA_UTAMA)
            sel.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for nomor_baris, baris in enumerate(data, start=1):
            warna = WARNA_LATAR if nomor_baris % 2 else WARNA_PERMUKAAN
            for indeks, nilai in enumerate(baris, start=1):
                sel = lembar.cell(row=baris_mulai + nomor_baris, column=indeks, value=nilai)
                self._isi_sel_xlsx(sel, warna_isi=warna)
        if freeze_pane:
            lembar.freeze_panes = freeze_pane
        self._atur_lebar_kolom_xlsx(lembar)

    def _tulis_xlsx(self, jalur: Path, dossier: DossierSindikat, bundel: Dict[str, Any], sidik: str) -> None:
        workbook = Workbook()
        dashboard = workbook.active
        dashboard.title = "Dashboard"
        dashboard.sheet_properties.tabColor = WARNA_UTAMA
        for kolom in ["A", "B", "C", "D", "E", "F"]:
            dashboard.column_dimensions[kolom].width = 18

        dashboard.merge_cells("A1:F2")
        sel_judul = dashboard["A1"]
        sel_judul.value = f"Dossier Analisa Kasus\n{dossier.judul_kasus}"
        self._isi_sel_xlsx(sel_judul, tebal=True, warna_teks=WARNA_LATAR, warna_isi=WARNA_UTAMA)
        sel_judul.font = Font(name=NAMA_FONT_UI, size=18, bold=True, color=WARNA_LATAR)
        sel_judul.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

        for rentang, label, nilai in [
            ("A4:B5", "ID Kasus", dossier.id_kasus),
            ("C4:D5", "Indikasi Sindikat", self._ya_tidak(dossier.indikasi_sindikat)),
            ("E4:F5", "Confidence", self._format_persentase(dossier.confidence)),
            ("A7:F8", "Fingerprint SHA-256", sidik),
        ]:
            dashboard.merge_cells(rentang)
            sel = dashboard[rentang.split(":")[0]]
            sel.value = f"{label}\n{nilai}"
            self._isi_sel_xlsx(sel, tebal=True, warna_isi=WARNA_PERMUKAAN)
            sel.border = Border(
                left=Side(style="medium", color=WARNA_UTAMA),
                right=Side(style="thin", color=WARNA_BATAS),
                top=Side(style="thin", color=WARNA_BATAS),
                bottom=Side(style="thin", color=WARNA_BATAS),
            )
            sel.alignment = Alignment(vertical="center", wrap_text=True)

        baris = self._tambah_judul_seksi_xlsx(dashboard, 10, "Ringkasan Eksekutif", "F")
        dashboard.merge_cells(f"A{baris}:F{baris + 2}")
        sel_ringkasan = dashboard[f"A{baris}"]
        sel_ringkasan.value = dossier.ringkasan_eksekutif
        self._isi_sel_xlsx(sel_ringkasan, warna_isi=WARNA_PERMUKAAN)
        baris += 4

        for judul, items in [
            ("Alasan Utama", dossier.alasan_utama),
            ("Pola Koordinasi", dossier.pola_koordinasi),
            ("Bukti Utama", dossier.bukti_utama),
            ("Bukti Lemah", dossier.bukti_lemah),
            ("Rekomendasi Lanjutan", dossier.rekomendasi_lanjutan),
        ]:
            baris = self._tambah_judul_seksi_xlsx(dashboard, baris, judul, "F")
            for item in items or ["Belum ada butir yang tersedia."]:
                dashboard.merge_cells(f"A{baris}:F{baris}")
                sel = dashboard[f"A{baris}"]
                sel.value = f"• {item}"
                self._isi_sel_xlsx(sel, warna_isi=WARNA_LATAR)
                baris += 1
            baris += 1

        statistik = [[kunci, nilai] for kunci, nilai in self._ringkas_bundel(bundel).items()]
        self._tulis_tabel_xlsx(dashboard, baris, ["Metrik", "Nilai"], statistik)

        lembar_aktor = workbook.create_sheet("Aktor Inti")
        lembar_aktor.sheet_properties.tabColor = WARNA_SEKUNDER
        self._tambah_judul_seksi_xlsx(lembar_aktor, 1, "Aktor Inti", "E")
        data_aktor = [
            [item.id_profil or "-", item.nama, item.peran, self._format_persentase(item.confidence), item.alasan]
            for item in dossier.aktor_inti
        ] or [["-", "Belum ada aktor inti", "-", "-", "-"]]
        self._tulis_tabel_xlsx(lembar_aktor, 3, ["ID Profil", "Nama", "Peran", "Confidence", "Alasan"], data_aktor, "A4")

        lembar_relasi = workbook.create_sheet("Relasi Kunci")
        lembar_relasi.sheet_properties.tabColor = WARNA_SEKUNDER
        self._tambah_judul_seksi_xlsx(lembar_relasi, 1, "Relasi Kunci", "E")
        data_relasi = [
            [item.sumber, item.target, item.jenis_relasi, self._format_persentase(item.confidence), item.alasan]
            for item in dossier.relasi_kunci
        ] or [["-", "-", "Belum ada relasi kunci", "-", "-"]]
        self._tulis_tabel_xlsx(lembar_relasi, 3, ["Sumber", "Target", "Jenis Relasi", "Confidence", "Alasan"], data_relasi, "A4")

        lembar_transaksi = workbook.create_sheet("Transaksi")
        lembar_transaksi.sheet_properties.tabColor = WARNA_SEKUNDER
        self._tambah_judul_seksi_xlsx(lembar_transaksi, 1, "Transaksi Terkait", "F")
        data_transaksi = [
            [
                item.get("id_transaksi", "-"),
                item.get("nama_sumber", "-"),
                item.get("nama_tujuan", "-"),
                self._format_uang_rupiah(item.get("jumlah_idr", 0)),
                item.get("kanal", "-"),
                self._format_waktu(item.get("timestamp")),
            ]
            for item in bundel.get("transaksi", [])
        ] or [["-", "-", "-", "-", "-", "-"]]
        self._tulis_tabel_xlsx(lembar_transaksi, 3, ["ID Transaksi", "Sumber", "Tujuan", "Jumlah", "Kanal", "Waktu"], data_transaksi, "A4")

        lembar_bukti = workbook.create_sheet("Bukti & Rekom")
        lembar_bukti.sheet_properties.tabColor = WARNA_UTAMA
        self._tambah_judul_seksi_xlsx(lembar_bukti, 1, "Bukti, Kelemahan, dan Rekomendasi", "C")
        jumlah_baris = max(len(dossier.bukti_utama), len(dossier.bukti_lemah), len(dossier.rekomendasi_lanjutan), 1)
        data_bukti = []
        for indeks in range(jumlah_baris):
            data_bukti.append(
                [
                    dossier.bukti_utama[indeks] if indeks < len(dossier.bukti_utama) else "",
                    dossier.bukti_lemah[indeks] if indeks < len(dossier.bukti_lemah) else "",
                    dossier.rekomendasi_lanjutan[indeks] if indeks < len(dossier.rekomendasi_lanjutan) else "",
                ]
            )
        self._tulis_tabel_xlsx(lembar_bukti, 3, ["Bukti Utama", "Bukti Lemah", "Rekomendasi Lanjutan"], data_bukti, "A4")

        workbook.save(jalur)

    def _atur_shading_sel_docx(self, sel, warna_hex: str) -> None:
        properti = sel._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), warna_hex)
        properti.append(shading)

    def _atur_margin_sel_docx(self, sel, ukuran: int = 120) -> None:
        properti = sel._tc.get_or_add_tcPr()
        margin = OxmlElement("w:tcMar")
        for sisi in ["top", "left", "bottom", "right"]:
            elemen = OxmlElement(f"w:{sisi}")
            elemen.set(qn("w:w"), str(ukuran))
            elemen.set(qn("w:type"), "dxa")
            margin.append(elemen)
        properti.append(margin)

    def _tambah_tabel_docx(self, dokumen: Document, header: List[str], data: List[List[Any]]) -> None:
        tabel = dokumen.add_table(rows=1, cols=len(header))
        tabel.style = "Table Grid"
        tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
        for indeks, judul in enumerate(header):
            sel = tabel.rows[0].cells[indeks]
            self._atur_shading_sel_docx(sel, WARNA_UTAMA)
            self._atur_margin_sel_docx(sel)
            run = sel.paragraphs[0].add_run(judul)
            run.font.name = NAMA_FONT_UI
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WARNA_LATAR)
            sel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for nomor, baris in enumerate(data, start=1):
            sel_baris = tabel.add_row().cells
            for indeks, nilai in enumerate(baris):
                sel = sel_baris[indeks]
                self._atur_shading_sel_docx(sel, WARNA_LATAR if nomor % 2 else WARNA_PERMUKAAN)
                self._atur_margin_sel_docx(sel)
                run = sel.paragraphs[0].add_run(str(nilai))
                run.font.name = NAMA_FONT_UI
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string(WARNA_TEKS)

    def _tambah_box_docx(self, dokumen: Document, judul: str, items: List[str], warna_judul: str = WARNA_UTAMA) -> None:
        tabel = dokumen.add_table(rows=1, cols=1)
        tabel.style = "Table Grid"
        sel = tabel.cell(0, 0)
        self._atur_shading_sel_docx(sel, WARNA_PERMUKAAN)
        self._atur_margin_sel_docx(sel, 140)
        p_judul = sel.paragraphs[0]
        run_judul = p_judul.add_run(judul)
        run_judul.font.name = NAMA_FONT_UI
        run_judul.font.size = Pt(11)
        run_judul.font.bold = True
        run_judul.font.color.rgb = RGBColor.from_string(warna_judul)
        for item in items or ["Belum ada butir yang tersedia."]:
            p_item = sel.add_paragraph(style="List Bullet")
            run_item = p_item.add_run(item)
            run_item.font.name = NAMA_FONT_UI
            run_item.font.size = Pt(10.5)
            run_item.font.color.rgb = RGBColor.from_string(WARNA_TEKS)

    def _tulis_docx(self, jalur: Path, dossier: DossierSindikat, bundel: Dict[str, Any], sidik: str) -> None:
        dokumen = Document()
        section = dokumen.sections[0]
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

        for nama_gaya, ukuran, warna in [("Normal", 10.5, WARNA_TEKS), ("Heading 1", 16, WARNA_UTAMA), ("Heading 2", 13, WARNA_TEKS), ("Heading 3", 11, WARNA_TEKS_REDUP)]:
            gaya = dokumen.styles[nama_gaya]
            gaya.font.name = NAMA_FONT_UI
            gaya.font.size = Pt(ukuran)
            gaya.font.bold = nama_gaya != "Normal"
            gaya.font.color.rgb = RGBColor.from_string(warna)

        banner = dokumen.add_table(rows=1, cols=1)
        banner.style = "Table Grid"
        sel_banner = banner.cell(0, 0)
        self._atur_shading_sel_docx(sel_banner, WARNA_UTAMA)
        self._atur_margin_sel_docx(sel_banner, 180)
        p_banner = sel_banner.paragraphs[0]
        run_judul = p_banner.add_run("DOSSIER ANALISA KASUS\n")
        run_judul.font.name = NAMA_FONT_UI
        run_judul.font.size = Pt(18)
        run_judul.font.bold = True
        run_judul.font.color.rgb = RGBColor.from_string(WARNA_LATAR)
        run_sub = p_banner.add_run(dossier.judul_kasus)
        run_sub.font.name = NAMA_FONT_UI
        run_sub.font.size = Pt(12)
        run_sub.font.bold = True
        run_sub.font.color.rgb = RGBColor.from_string(WARNA_LATAR)

        dokumen.add_paragraph()
        meta = dokumen.add_table(rows=2, cols=2)
        meta.style = "Table Grid"
        for sel, label, nilai in [
            (meta.cell(0, 0), "ID Kasus", dossier.id_kasus),
            (meta.cell(0, 1), "Indikasi Sindikat", self._ya_tidak(dossier.indikasi_sindikat)),
            (meta.cell(1, 0), "Confidence", self._format_persentase(dossier.confidence)),
            (meta.cell(1, 1), "Fingerprint SHA-256", sidik),
        ]:
            self._atur_shading_sel_docx(sel, WARNA_PERMUKAAN)
            self._atur_margin_sel_docx(sel)
            p = sel.paragraphs[0]
            run_label = p.add_run(f"{label}\n")
            run_label.font.name = NAMA_FONT_UI
            run_label.font.size = Pt(9)
            run_label.font.bold = True
            run_label.font.color.rgb = RGBColor.from_string(WARNA_TEKS_REDUP)
            run_nilai = p.add_run(nilai)
            run_nilai.font.name = NAMA_FONT_UI
            run_nilai.font.size = Pt(12)
            run_nilai.font.bold = True
            run_nilai.font.color.rgb = RGBColor.from_string(WARNA_TEKS)

        dokumen.add_heading("Ringkasan Eksekutif", level=1)
        self._tambah_box_docx(dokumen, "Intisari", [dossier.ringkasan_eksekutif])
        dokumen.add_heading("Temuan Inti", level=1)
        dokumen.add_heading("Alasan Utama", level=2)
        self._tambah_box_docx(dokumen, "Dasar Penilaian", dossier.alasan_utama)
        dokumen.add_heading("Pola Koordinasi", level=2)
        self._tambah_box_docx(dokumen, "Pola yang Terdeteksi", dossier.pola_koordinasi, WARNA_PERINGATAN)
        dokumen.add_heading("Entitas dan Relasi", level=1)
        dokumen.add_heading("Aktor Inti", level=2)
        dokumen.add_heading("Tabel Aktor Prioritas", level=3)
        data_aktor = [[item.id_profil or "-", item.nama, item.peran, self._format_persentase(item.confidence), item.alasan] for item in dossier.aktor_inti] or [["-", "Belum ada aktor inti", "-", "-", "-"]]
        self._tambah_tabel_docx(dokumen, ["ID Profil", "Nama", "Peran", "Confidence", "Alasan"], data_aktor)
        dokumen.add_heading("Relasi Kunci", level=2)
        dokumen.add_heading("Tabel Relasi Tervalidasi", level=3)
        data_relasi = [[item.sumber, item.target, item.jenis_relasi, self._format_persentase(item.confidence), item.alasan] for item in dossier.relasi_kunci] or [["-", "-", "Belum ada relasi kunci", "-", "-"]]
        self._tambah_tabel_docx(dokumen, ["Sumber", "Target", "Jenis Relasi", "Confidence", "Alasan"], data_relasi)
        dokumen.add_heading("Bukti dan Rekomendasi", level=1)
        dokumen.add_heading("Bukti Utama", level=2)
        self._tambah_box_docx(dokumen, "Bukti Penguat", dossier.bukti_utama, WARNA_SUKSES)
        dokumen.add_heading("Bukti Lemah", level=2)
        self._tambah_box_docx(dokumen, "Area yang Masih Lemah", dossier.bukti_lemah, WARNA_BAHAYA)
        dokumen.add_heading("Rekomendasi Lanjutan", level=2)
        self._tambah_box_docx(dokumen, "Tindak Lanjut", dossier.rekomendasi_lanjutan, WARNA_UTAMA)
        dokumen.add_heading("Narasi Analisis", level=1)
        dokumen.add_paragraph(dossier.narasi_analisis or "Narasi analisis belum tersedia.")
        dokumen.add_heading("Statistik Bundel", level=1)
        dokumen.add_heading("Ringkasan Data Sumber", level=3)
        self._tambah_tabel_docx(dokumen, ["Metrik", "Nilai"], [[k, v] for k, v in self._ringkas_bundel(bundel).items()])
        dokumen.save(jalur)

    def _font_pdf(self) -> tuple[str, str]:
        kandidat = [
            ("SegoeUI", "SegoeUI-Bold", Path("C:/Windows/Fonts/segoeui.ttf"), Path("C:/Windows/Fonts/segoeuib.ttf")),
            ("Aptos", "Aptos-Bold", Path("C:/Windows/Fonts/aptos.ttf"), Path("C:/Windows/Fonts/aptos-bold.ttf")),
        ]
        for nama_reguler, nama_tebal, jalur_reguler, jalur_tebal in kandidat:
            if jalur_reguler.exists() and jalur_tebal.exists():
                if nama_reguler not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(nama_reguler, str(jalur_reguler)))
                if nama_tebal not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(nama_tebal, str(jalur_tebal)))
                return nama_reguler, nama_tebal
        return "Helvetica", "Helvetica-Bold"

    def _gaya_pdf(self) -> dict[str, ParagraphStyle]:
        font_reguler, font_tebal = self._font_pdf()
        sampel = getSampleStyleSheet()
        return {
            "judul": ParagraphStyle("judul", parent=sampel["Heading1"], fontName=font_tebal, fontSize=20, leading=24, textColor=colors.white),
            "subjudul": ParagraphStyle("subjudul", parent=sampel["BodyText"], fontName=font_tebal, fontSize=11, leading=14, textColor=colors.white),
            "h1": ParagraphStyle("h1", parent=sampel["Heading1"], fontName=font_tebal, fontSize=15, leading=18, textColor=colors.HexColor(f"#{WARNA_UTAMA}"), spaceBefore=8, spaceAfter=6),
            "h2": ParagraphStyle("h2", parent=sampel["Heading2"], fontName=font_tebal, fontSize=12, leading=14, textColor=colors.HexColor(f"#{WARNA_TEKS}"), spaceBefore=5, spaceAfter=4),
            "h3": ParagraphStyle("h3", parent=sampel["Heading3"], fontName=font_tebal, fontSize=10, leading=12, textColor=colors.HexColor(f"#{WARNA_TEKS_REDUP}"), spaceBefore=4, spaceAfter=3),
            "isi": ParagraphStyle("isi", parent=sampel["BodyText"], fontName=font_reguler, fontSize=10, leading=14, textColor=colors.HexColor(f"#{WARNA_TEKS}")),
            "label": ParagraphStyle("label", parent=sampel["BodyText"], fontName=font_tebal, fontSize=8, leading=10, alignment=TA_CENTER, textColor=colors.HexColor(f"#{WARNA_TEKS_REDUP}")),
            "nilai": ParagraphStyle("nilai", parent=sampel["BodyText"], fontName=font_tebal, fontSize=12, leading=14, alignment=TA_CENTER, textColor=colors.HexColor(f"#{WARNA_TEKS}")),
        }

    def _kotak_pdf(self, judul: str, items: List[str], gaya: dict[str, ParagraphStyle], warna_judul: str) -> Table:
        data = [[Paragraph(f"<b>{escape_xml(judul)}</b>", gaya["h3"])]]
        for item in items or ["Belum ada butir yang tersedia."]:
            data.append([Paragraph(f"• {escape_xml(str(item))}", gaya["isi"])])
        tabel = Table(data, colWidths=[180 * mm])
        tabel.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{WARNA_PERMUKAAN}")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{warna_judul}")),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor(f"#{WARNA_LATAR}")),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor(f"#{WARNA_BATAS}")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{WARNA_BATAS}")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        return tabel

    def _tabel_pdf(self, header: List[str], data: List[List[Any]], gaya: dict[str, ParagraphStyle], lebar: List[float]) -> Table:
        konten = [[Paragraph(f"<b>{escape_xml(str(item))}</b>", gaya["isi"]) for item in header]]
        for baris in data:
            konten.append([Paragraph(escape_xml(str(item)), gaya["isi"]) for item in baris])
        tabel = Table(konten, colWidths=lebar, repeatRows=1)
        tabel.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{WARNA_UTAMA}")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor(f"#{WARNA_LATAR}")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(f"#{WARNA_LATAR}"), colors.HexColor(f"#{WARNA_PERMUKAAN}")]),
                    ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor(f"#{WARNA_BATAS}")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        return tabel

    def _hias_halaman_pdf(self, kanvas, dokumen) -> None:
        kanvas.saveState()
        lebar, tinggi = A4
        kanvas.setFillColor(colors.HexColor(f"#{WARNA_UTAMA}"))
        kanvas.rect(0, tinggi - 9 * mm, lebar, 9 * mm, stroke=0, fill=1)
        kanvas.setFillColor(colors.HexColor(f"#{WARNA_TEKS_REDUP}"))
        kanvas.setFont(self._font_pdf()[0], 8)
        kanvas.drawRightString(lebar - 15 * mm, 10 * mm, f"Halaman {dokumen.page}")
        kanvas.restoreState()

    def _tulis_pdf(self, jalur: Path, dossier: DossierSindikat, bundel: Dict[str, Any], sidik: str) -> None:
        gaya = self._gaya_pdf()
        elemen = []

        banner = Table([[Paragraph("DOSSIER ANALISA KASUS", gaya["judul"]), Paragraph(escape_xml(dossier.judul_kasus), gaya["subjudul"])]], colWidths=[180 * mm])
        banner.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{WARNA_UTAMA}")), ("LEFTPADDING", (0, 0), (-1, -1), 14), ("RIGHTPADDING", (0, 0), (-1, -1), 14), ("TOPPADDING", (0, 0), (-1, -1), 12), ("BOTTOMPADDING", (0, 0), (-1, -1), 12)]))
        elemen.extend([banner, Spacer(1, 8)])

        kartu = Table([[Paragraph("ID KASUS", gaya["label"]), Paragraph("INDIKASI", gaya["label"]), Paragraph("CONFIDENCE", gaya["label"])], [Paragraph(escape_xml(dossier.id_kasus), gaya["nilai"]), Paragraph(self._ya_tidak(dossier.indikasi_sindikat), gaya["nilai"]), Paragraph(self._format_persentase(dossier.confidence), gaya["nilai"])]], colWidths=[60 * mm, 60 * mm, 60 * mm])
        kartu.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{WARNA_PERMUKAAN}")), ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor(f"#{WARNA_LATAR}")), ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor(f"#{WARNA_BATAS}")), ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor(f"#{WARNA_BATAS}")), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6)]))
        elemen.extend([kartu, Spacer(1, 10), Paragraph("Ringkasan Eksekutif", gaya["h1"]), self._kotak_pdf("Intisari", [dossier.ringkasan_eksekutif], gaya, WARNA_UTAMA), Spacer(1, 8)])

        elemen.extend([Paragraph("Temuan Inti", gaya["h1"]), Paragraph("Alasan Utama", gaya["h2"]), self._kotak_pdf("Dasar Penilaian", dossier.alasan_utama, gaya, WARNA_UTAMA), Spacer(1, 6), Paragraph("Pola Koordinasi", gaya["h2"]), self._kotak_pdf("Pola yang Terdeteksi", dossier.pola_koordinasi, gaya, WARNA_PERINGATAN), Spacer(1, 8)])

        data_aktor = [[item.id_profil or "-", item.nama, item.peran, self._format_persentase(item.confidence), item.alasan] for item in dossier.aktor_inti] or [["-", "Belum ada aktor inti", "-", "-", "-"]]
        data_relasi = [[item.sumber, item.target, item.jenis_relasi, self._format_persentase(item.confidence), item.alasan] for item in dossier.relasi_kunci] or [["-", "-", "Belum ada relasi kunci", "-", "-"]]
        elemen.extend([Paragraph("Entitas dan Relasi", gaya["h1"]), Paragraph("Aktor Inti", gaya["h2"]), Paragraph("Tabel Aktor Prioritas", gaya["h3"]), self._tabel_pdf(["ID Profil", "Nama", "Peran", "Confidence", "Alasan"], data_aktor, gaya, [24 * mm, 30 * mm, 32 * mm, 22 * mm, 72 * mm]), Spacer(1, 8), Paragraph("Relasi Kunci", gaya["h2"]), Paragraph("Tabel Relasi Tervalidasi", gaya["h3"]), self._tabel_pdf(["Sumber", "Target", "Jenis Relasi", "Confidence", "Alasan"], data_relasi, gaya, [28 * mm, 28 * mm, 30 * mm, 22 * mm, 62 * mm]), Spacer(1, 8)])

        elemen.extend([Paragraph("Bukti dan Rekomendasi", gaya["h1"]), Paragraph("Bukti Utama", gaya["h2"]), self._kotak_pdf("Bukti Penguat", dossier.bukti_utama, gaya, WARNA_SUKSES), Spacer(1, 6), Paragraph("Bukti Lemah", gaya["h2"]), self._kotak_pdf("Area yang Masih Lemah", dossier.bukti_lemah, gaya, WARNA_BAHAYA), Spacer(1, 6), Paragraph("Rekomendasi Lanjutan", gaya["h2"]), self._kotak_pdf("Tindak Lanjut", dossier.rekomendasi_lanjutan, gaya, WARNA_UTAMA), Spacer(1, 8)])

        elemen.extend([Paragraph("Narasi Analisis", gaya["h1"]), Paragraph(escape_xml(dossier.narasi_analisis or "Narasi analisis belum tersedia."), gaya["isi"]), Spacer(1, 8), Paragraph("Statistik Bundel", gaya["h1"]), Paragraph("Ringkasan Data Sumber", gaya["h3"]), self._tabel_pdf(["Metrik", "Nilai"], [[k, v] for k, v in self._ringkas_bundel(bundel).items()], gaya, [90 * mm, 90 * mm]), Spacer(1, 8), Paragraph(f"Fingerprint SHA-256: {escape_xml(sidik)}", gaya["isi"])])

        dokumen = SimpleDocTemplate(str(jalur), pagesize=A4, leftMargin=15 * mm, rightMargin=15 * mm, topMargin=18 * mm, bottomMargin=15 * mm)
        dokumen.build(elemen, onFirstPage=self._hias_halaman_pdf, onLaterPages=self._hias_halaman_pdf)


class MesinAnalisisKasus:
    def __init__(
        self,
        pengaturan: PengaturanRuntime,
        mcp: McpGateway,
        layanan_openai: LayananOpenAI,
    ):
        self.pengaturan = pengaturan
        self.mcp = mcp
        self.layanan_openai = layanan_openai
        self.penghasil = PenghasilArtefakAnalisa((pengaturan.akar_data / "analisa").resolve())

    def analisis(self, id_kasus: str) -> Dict[str, Any]:
        trace_id = f"trace-kasus-{_sha256_teks(id_kasus)[:12]}"
        logger = logger_dengan_trace(__name__, trace_id)
        print(f"[ANALISIS] Mengambil bundel kasus `{id_kasus}` dari PostgreSQL dan Neo4j...")
        bundel = self.mcp.ambil_bundle_kasus(trace_id, id_kasus)
        logger.info(
            "Bundel kasus berhasil diambil",
            extra={
                "extra_payload": {
                    "id_kasus": id_kasus,
                    "jumlah_transaksi": len(bundel.get("transaksi", [])),
                    "jumlah_profil": len(bundel.get("profil", [])),
                    "jumlah_laporan": len(bundel.get("laporan", [])),
                    "jumlah_node_graf": len((bundel.get("graf") or {}).get("nodes", [])),
                }
            },
        )
        print("[ANALISIS] Bundel kasus siap. Mengirim reasoning ke OpenAI...")
        dossier = self.layanan_openai.analisis_kasus_sindikat(
            kasus=bundel["kasus"],
            laporan=bundel["laporan"],
            skor_risiko=bundel["skor_risiko"],
            transaksi=bundel["transaksi"],
            kampanye=bundel["kampanye"],
            profil=bundel["profil"],
            lokasi=bundel["lokasi"],
            postingan=bundel["postingan"],
            graf=bundel["graf"],
        )
        logger.info(
            "Dossier sindikat berhasil diterima dari OpenAI",
            extra={
                "extra_payload": {
                    "id_kasus": id_kasus,
                    "indikasi_sindikat": dossier.indikasi_sindikat,
                    "confidence": dossier.confidence,
                    "jumlah_aktor_inti": len(dossier.aktor_inti),
                    "jumlah_relasi_kunci": len(dossier.relasi_kunci),
                }
            },
        )
        print("[ANALISIS] Dossier selesai. Menulis artefak ke folder analisa...")
        hasil = self.penghasil.simpan_semua(id_kasus, dossier, bundel)
        logger.info(
            "Analisis kasus selesai dan artefak berhasil dibuat",
            extra={
                "extra_payload": {
                    "id_kasus": id_kasus,
                    "direktori": hasil["direktori"],
                    "fingerprint_sha256": hasil["fingerprint_sha256"],
                }
            },
        )
        print(f"[ANALISIS] Artefak berhasil dibuat di: {hasil['direktori']}")
        print(f"[ANALISIS] Fingerprint SHA-256: {hasil['fingerprint_sha256']}")
        return {
            "trace_id": trace_id,
            "dossier": dossier.model_dump(mode="json"),
            "artefak": hasil,
        }
